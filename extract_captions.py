import os
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

def extract_tables_with_captions(source_path, dest_path):
    src_doc = Document(source_path)
    supp_doc = Document()
    
    supp_doc.add_heading('Supplemental Material: Tables and Descriptions', 0)
    
    body = src_doc.element.body
    elements = list(body)
    
    count = 1
    for i, element in enumerate(elements):
        if isinstance(element, CT_Tbl):
            # Try to get previous paragraph
            prev_p = None
            for j in range(i-1, -1, -1):
                if isinstance(elements[j], CT_P):
                    p = Paragraph(elements[j], src_doc)
                    if p.text.strip():
                        prev_p = p.text.strip()
                        break
            
            # Try to get next paragraph
            next_p = None
            for j in range(i+1, len(elements)):
                if isinstance(elements[j], CT_P):
                    p = Paragraph(elements[j], src_doc)
                    if p.text.strip():
                        next_p = p.text.strip()
                        break

            # Write caption if it looks like one
            caption_added = False
            if prev_p and ('table' in prev_p.lower() or 'tabel' in prev_p.lower()):
                supp_doc.add_paragraph(prev_p, style='Caption')
                caption_added = True
            elif next_p and ('table' in next_p.lower() or 'tabel' in next_p.lower()):
                supp_doc.add_paragraph(next_p, style='Caption')
                caption_added = True
            
            if not caption_added:
                supp_doc.add_paragraph(f"Supplemental Table {count}", style='Caption')
                if prev_p and len(prev_p) > 20: # arbitrary length for a description
                     supp_doc.add_paragraph(prev_p)

            # Copy the table
            table = Table(element, src_doc)
            new_table = supp_doc.add_table(rows=0, cols=len(table.columns))
            new_table.style = 'Table Grid'
            
            for row in table.rows:
                new_row = new_table.add_row()
                for idx, cell in enumerate(row.cells):
                    try:
                        new_row.cells[idx].text = cell.text
                    except IndexError:
                        pass # handle merged cells roughly
            
            supp_doc.add_paragraph("") # Space
            count += 1
            
    supp_doc.save(dest_path)
    print(f"Extracted tables and captions to {dest_path}")

if __name__ == '__main__':
    extract_tables_with_captions('Paper/jama/brain_ctc_seg-3.docx', 'Paper/jama/Supplemental_Material.docx')
