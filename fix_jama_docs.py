import os
import re
from docx import Document
from docx.shared import Pt
import copy

def process_equations(doc_path):
    doc = Document(doc_path)
    
    replacements = {
        r'\\text\{Dilate\}': 'Dilate',
        r'\\text\{Erode\}': 'Erode',
        r'\\text\{CE\}': 'CE',
        r'\\text\{IoU\}': 'IoU',
        r'\\mathcal\{L\}': 'L',
        r'\\mathbb\{R\}\^2': 'R²',
        r'\\theta': 'θ',
        r'\\lambda': 'λ',
        r'\\odot': '⊙',
        r'\\cdot': '·',
        r'\\cap': '∩',
        r'\\cup': '∪',
        r'\\hat\{Y\}': 'Y_hat',
        r'\\frac\{(.*?)\}\{(.*?)\}': r'\1 / \2',
        r'\$': '' # Remove inline math delimiters
    }

    def format_run_text(p, text):
        # Clear existing runs
        p.clear()
        
        # Apply regex replacements
        for k, v in replacements.items():
            text = re.sub(k, v, text)
            
        # Parse for subscripts and superscripts manually
        # Simple parser for A_b or A^b
        tokens = re.split(r'([_^]\{[^}]+\}|[_^].)', text)
        for token in tokens:
            if not token: continue
            if token.startswith('_'):
                sub_text = token[1:].strip('{}')
                run = p.add_run(sub_text)
                run.font.subscript = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            elif token.startswith('^'):
                sup_text = token[1:].strip('{}')
                run = p.add_run(sup_text)
                run.font.superscript = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            else:
                run = p.add_run(token)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    for p in doc.paragraphs:
        if '\\' in p.text or '_' in p.text or '^' in p.text or '$' in p.text:
            if 'Corresponding' in p.text or 'Word count' in p.text:
                continue
            original_text = p.text
            format_run_text(p, original_text)
            
    doc.save('Paper/jama/Manuscript_JAMA_Final2_Fixed.docx')
    print("Fixed equations in Manuscript_JAMA_Final2_Fixed.docx")

def extract_tables_to_supplemental(source_path, dest_path):
    src_doc = Document(source_path)
    supp_doc = Document()
    
    supp_doc.add_heading('Supplemental Material: Result Tables', 0)
    
    count = 1
    for table in src_doc.tables:
        supp_doc.add_paragraph(f"Supplemental Table {count}", style='Caption')
        
        # Create a new table in the supplemental doc
        new_table = supp_doc.add_table(rows=0, cols=len(table.columns))
        new_table.style = 'Table Grid'
        
        for row in table.rows:
            new_row = new_table.add_row()
            for idx, cell in enumerate(row.cells):
                new_row.cells[idx].text = cell.text
        
        supp_doc.add_paragraph("") # Space
        count += 1
        
    supp_doc.save(dest_path)
    print(f"Extracted {count-1} tables to {dest_path}")

if __name__ == '__main__':
    process_equations('Paper/jama/Manuscript_JAMA_Final2.docx')
    extract_tables_to_supplemental('Paper/jama/brain_ctc_seg-3.docx', 'Paper/jama/Supplemental_Material.docx')
