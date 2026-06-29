import os
from docx import Document
from docx.shared import Inches

def append_figures_to_supplemental(supp_path, media_dir):
    doc = Document(supp_path)
    
    doc.add_page_break()
    doc.add_heading('Supplemental Material: Figures', 0)
    
    # Sort images to maintain some order (by number in filename)
    import re
    def extract_number(f):
        match = re.search(r'\d+', f)
        return int(match.group()) if match else 0
        
    image_files = sorted([f for f in os.listdir(media_dir) if f.endswith(('.png', '.jpg', '.jpeg'))], key=extract_number)
    
    count = 1
    for img_file in image_files:
        img_path = os.path.join(media_dir, img_file)
        
        doc.add_paragraph(f"Supplemental Figure {count}", style='Caption')
        try:
            doc.add_picture(img_path, width=Inches(6.0))
        except Exception as e:
            print(f"Error adding {img_file}: {e}")
            
        doc.add_paragraph("") # Space
        count += 1
        
    doc.save(supp_path)
    print(f"Appended {count-1} figures to {supp_path}")

if __name__ == '__main__':
    append_figures_to_supplemental('Paper/jama/Supplemental_Material.docx', 'Paper/jama/word/media')
