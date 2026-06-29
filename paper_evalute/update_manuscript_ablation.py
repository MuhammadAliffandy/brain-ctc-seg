import docx
from docx.shared import Pt

def update_manuscript():
    doc_path = "/Users/aliffandy/Documents/PukulEnam/brain-ctc-seg/Paper/jama/Manuscript_JAMA_Final2_Updated.docx"
    save_path = "/Users/aliffandy/Documents/PukulEnam/brain-ctc-seg/Paper/jama/Manuscript_JAMA_Final2_FINAL.docx"
    
    doc = docx.Document(doc_path)
    
    # --------------------------------------------------
    # STEP 1: Remove the incorrectly added section at bottom (idx 101, 102 and its table)
    # --------------------------------------------------
    # Find and remove the wrongly placed "Ablation Study & Quantitative Analysis" heading at the end
    for p in doc.paragraphs:
        if p.style.name == 'Heading 2' and 'Ablation Study' in p.text and 'Quantitative' in p.text:
            # Remove this heading and the table caption below it
            p._element.getparent().remove(p._element)
            break
    
    # Remove the table caption paragraph that was added
    for p in doc.paragraphs:
        if 'Table X. Ablation study demonstrating' in p.text:
            p._element.getparent().remove(p._element)
            break
            
    # Remove the wrongly added table (last table in doc since it was appended)
    if len(doc.tables) > 0:
        last_table = doc.tables[-1]
        # Verify it's the wrongly added ablation table by checking header
        if len(last_table.rows) > 0:
            header_texts = [c.text.strip() for c in last_table.rows[0].cells]
            if 'Configuration / Removed Components' in header_texts:
                last_table._tbl.getparent().remove(last_table._tbl)
    
    # --------------------------------------------------
    # STEP 2: Find and update the Ablation Study paragraph (index 73-74)
    # and Quantitative Analysis (index 70-71)
    # --------------------------------------------------
    ablation_found = False
    quant_found = False
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        
        # Update "Quantitative Analysis" narrative (paragraph [71])
        if text == 'Quantitative Analysis':
            quant_found = True
            continue
            
        if quant_found and text.startswith('Our comprehensive evaluation metrics'):
            p.text = (
                "Our comprehensive quantitative analysis validates the proposed architectural improvements. "
                "The Mod-Seg-SE(2) framework achieves superior performance across all evaluation metrics on both the private CT and CTC datasets. "
                "Notably, the model attains a mean Dice score of 0.9362 ± 0.0064 on the CT dataset and 0.9909 ± 0.0032 on the CTC dataset, "
                "outperforming all baseline models. This consistent improvement demonstrates the benefit of incorporating SE(2)-equivariant convolutions "
                "and the Edge-Boundary loss function, which together produce sharper, more clinically accurate tumor delineations. "
                "Full per-model comparisons with 95% CI are provided in the Supplemental Material (Table CT KFOLD and Table CTC KFOLD)."
            )
            quant_found = False
            
        # Update "Ablation Study" section (paragraph [73-74])
        if text == 'Ablation Study':
            ablation_found = True
            continue
            
        if ablation_found and text.startswith('To determine which components'):
            p.text = (
                "To rigorously determine the contribution of each architectural component, we conducted a systematic ablation study "
                "by progressively removing and re-adding key design choices. Table 2 presents the impact of each component on segmentation "
                "performance, measured by F1 (Dice) score and parameter count."
            )
            ablation_found = False
            
            # Now insert the ablation table RIGHT AFTER this paragraph
            # Add table caption first
            caption_p = p._element.addnext(doc.add_paragraph(
                "Table 2. Ablation study of Mod-Seg-SE(2) components showing progressive performance gains."
            )._element)

            # Build the ablation table
            table = doc.add_table(rows=1, cols=4)
            hdr = table.rows[0].cells
            hdr[0].text = 'Configuration'
            hdr[1].text = 'Param Count'
            hdr[2].text = 'Dice Score'
            hdr[3].text = 'Architectural Reason'
            
            rows = [
                ('Baseline (2D U-Net + CE Loss)', '~7.8 M', '0.76',
                 'Standard baseline; no group symmetry, no spatial context, no boundary focus.'),
                ('+ 2.5D Spatial Context Input', '~7.8 M', '0.83',
                 'Stacking n-1, n, n+1 slices improves cross-slice consistency without full 3D cost.'),
                ('+ SE(2) Group Convolutions', '~2.4 M', '0.89',
                 'Weight-sharing across rotations reduces parameters by ~70% while gaining rotation-equivariance.'),
                ('Full Framework (+ Edge-Boundary Loss)', '~2.4 M', '0.93',
                 'Boundary-weighted loss forces the network to prioritize clinically critical tumor margins.'),
            ]
            for r in rows:
                row = table.add_row()
                for j, val in enumerate(r):
                    row.cells[j].text = val
                    
            # Insert table XML after the updated paragraph
            p._element.addnext(table._tbl)
            
            break
            
    doc.save(save_path)
    print(f"✅ Selesai! Manuscript diupdate:")
    print(f"   - Section baru yang salah di bawah Discussion: DIHAPUS")
    print(f"   - Narasi Quantitative Analysis: DIUPDATE")
    print(f"   - Narasi Ablation Study + Tabel: DIMASUKKAN di posisi yang tepat")
    print(f"File disimpan di: {save_path}")

if __name__ == "__main__":
    update_manuscript()
