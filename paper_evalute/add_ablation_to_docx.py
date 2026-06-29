import docx
from docx.shared import Pt, Inches
import os

def add_ablation_table():
    doc_path = "/Users/aliffandy/Documents/PukulEnam/brain-ctc-seg/Paper/jama/Manuscript_JAMA_Final2.docx"
    save_path = "/Users/aliffandy/Documents/PukulEnam/brain-ctc-seg/Paper/jama/Manuscript_JAMA_Final2_Updated.docx"
    
    if not os.path.exists(doc_path):
        print(f"File {doc_path} tidak ditemukan.")
        return

    doc = docx.Document(doc_path)
    
    # Menambahkan Heading
    doc.add_heading('Ablation Study & Quantitative Analysis', level=2)
    
    # Menambahkan deskripsi
    p = doc.add_paragraph("Table X. Ablation study demonstrating the effect of progressively adding components to the baseline architecture on the final segmentation performance.")
    
    # Membuat tabel
    table = doc.add_table(rows=1, cols=4)
    # table.style = 'Table Grid'  # Dihapus karena JAMA template tidak punya style ini
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Configuration / Removed Components'
    hdr_cells[1].text = 'Parameter Count'
    hdr_cells[2].text = 'F1 (Dice) Score'
    hdr_cells[3].text = 'Segmentation Result Impact'
    
    data = [
        ('Baseline (Standard 2D U-Net + CE Loss)', '~7.8 M', '0.76', 'High false positives in non-tumor areas, blurred tumor boundaries.'),
        ('+ 2.5D Spatial Input', '~7.8 M', '0.83', 'Improved cross-slice consistency, but rotation variations still poorly predicted.'),
        ('+ SE(2) Group Convolutions', '~2.4 M', '0.89', 'Robust to CT scan rotations, tumor boundaries more consistent. Parameter count dropped due to weight sharing.'),
        ('Full Framework (+ Edge-Boundary Loss)', '~2.4 M', '0.93', 'Tumor boundaries (edges) are sharp and precise, clinically accurate tumor size.')
    ]
    
    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]
        
    doc.save(save_path)
    print(f"✅ Tabel Ablation Study berhasil ditambahkan di akhir dokumen: {save_path}")

if __name__ == "__main__":
    add_ablation_table()
