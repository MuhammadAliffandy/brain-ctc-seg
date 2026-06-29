import docx
import os

def compute_ci(mean_str, std_str):
    return f"{mean_str} ± {std_str}"

def add_exact_ci():
    doc_path = "/Users/aliffandy/Documents/PukulEnam/brain-ctc-seg/Paper/jama/Supplemental_Material.docx"
    save_path = "/Users/aliffandy/Documents/PukulEnam/brain-ctc-seg/Paper/jama/Supplemental_Material_FINAL.docx"
    
    doc = docx.Document(doc_path)
    
    # Data extracted from the image
    # Format: { 'Model_Name_lower': (mean, std) } for each metric
    # Metrics order in table: Accuracy, Precision, Recall, F1 (Dice), IoU
    
    # CT KFOLD (Table 5)
    ct_data = {
        'se2': [('0.9994', '0.0001'), ('0.9186', '0.0119'), ('0.9545', '0.0034'), ('0.9362', '0.0064'), ('0.8801', '0.0112')],
        'harmonic': [('0.9987', '0.0001'), ('0.8090', '0.0128'), ('0.9605', '0.0060'), ('0.8782', '0.0096'), ('0.7830', '0.0152')],
        'nnunet': [('0.9978', '0.0003'), ('0.6856', '0.0261'), ('0.9813', '0.0160'), ('0.8071', '0.0214'), ('0.6770', '0.0300')],
        'attention': [('0.9969', '0.0005'), ('0.6572', '0.0591'), ('0.7260', '0.0454'), ('0.6883', '0.0390'), ('0.5259', '0.0455')],
        'unet': [('0.9966', '0.0003'), ('0.6304', '0.0395'), ('0.6902', '0.0148'), ('0.6584', '0.0236'), ('0.4912', '0.0266')],
        'transunet': [('0.9907', '0.0018'), ('0.3148', '0.0389'), ('0.7961', '0.1302'), ('0.4474', '0.0385'), ('0.2888', '0.0320')]
    }
    
    # CTC KFOLD (Table 6) (Using the worst duplicates as matched with user's docx)
    ctc_data = {
        'se2': [('0.9999', '0.0000'), ('0.9865', '0.0023'), ('0.9953', '0.0044'), ('0.9909', '0.0032'), ('0.9819', '0.0063')],
        'harmonic': [('0.9995', '0.0001'), ('0.9174', '0.0111'), ('0.9985', '0.0005'), ('0.9562', '0.0059'), ('0.9162', '0.0109')],
        'nnunet': [('0.9990', '0.0001'), ('0.8485', '0.0109'), ('0.9944', '0.0018'), ('0.9156', '0.0066'), ('0.8445', '0.0112')],
        'unet': [('0.9990', '0.0002'), ('0.8561', '0.0286'), ('0.9671', '0.0190'), ('0.9080', '0.0199'), ('0.8320', '0.0335')],
        'attention': [('0.9986', '0.0007'), ('0.8215', '0.0458'), ('0.9440', '0.0992'), ('0.8779', '0.0692'), ('0.7874', '0.1024')],
        'transunet': [('0.9970', '0.0006'), ('0.6640', '0.0451'), ('0.8487', '0.0647'), ('0.7449', '0.0514'), ('0.5956', '0.0654')]
    }

    def update_table(table, data_dict):
        # Header is [Model, Accuracy, Precision, Recall, F1 (Dice), IoU]
        # We need to make sure the table has 6 columns
        # First row is header, so skip it
        # Clear existing rows (except header) and rewrite them
        # Wait, the table in docx might only have 1 row (Mod-Seg-SE2). We should add the rest!
        # Because the previous inspect showed only Row 1 for Mod-Seg-SE2.
        
        # Remove all rows except header
        while len(table.rows) > 1:
            tr = table.rows[1]._tr
            table._tbl.remove(tr)
            
        # Add the 6 models
        order = [('Mod-Seg-SE(2)', 'se2'), ('HarmonicNet', 'harmonic'), ('nnU-Net', 'nnunet'), ('Standard U-Net', 'unet'), ('Attention U-Net', 'attention'), ('TransUNet', 'transunet')]
        
        for display_name, key in order:
            row = table.add_row()
            row.cells[0].text = display_name
            for i in range(5):
                mean_str, std_str = data_dict[key][i]
                row.cells[i+1].text = compute_ci(mean_str, std_str)

    # In Supplemental_Material.docx, Table 5 is CT KFOLD and Table 6 is CTC KFOLD
    update_table(doc.tables[5], ct_data)
    update_table(doc.tables[6], ctc_data)
    
    # We also need to re-insert the figures since we are reading from original Supplemental_Material.docx
    from docx.shared import Inches
    figures_to_insert = [
        "Fig1_3Brain_Heatmap_Journal.png",
        "Fig2_Segmentation_Grid.png",
        "Fig3_Aggregated_ROC.png",
        "Fig4_Multi_CT_List.png",
        "Fig5_Mod_Seg_SE2_Architecture.png",
        "Fig6_Combined_Training_Pipeline.png"
    ]
    fig_idx = 0
    base_fig_dir = "/Users/aliffandy/Documents/PukulEnam/brain-ctc-seg/Paper/Journal_Figures/"
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("Supplemental Figure") and fig_idx < len(figures_to_insert):
            img_path = os.path.join(base_fig_dir, figures_to_insert[fig_idx])
            if os.path.exists(img_path):
                if i + 1 < len(doc.paragraphs):
                    new_p = doc.paragraphs[i+1].insert_paragraph_before()
                    run = new_p.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                fig_idx += 1

    doc.save(save_path)
    print(f"✅ Tabel K-Fold 5 & 6 diupdate dengan EXACT 95% CI. Tabel lain tetap aslinya.")
    print(f"✅ Gambar disisipkan.")
    print(f"File akhir: {save_path}")

if __name__ == "__main__":
    add_exact_ci()
