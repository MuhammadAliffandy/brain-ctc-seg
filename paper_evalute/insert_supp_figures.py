import docx
from docx.shared import Inches
import os

def insert_figures_to_supplemental():
    doc_path = "/Users/aliffandy/Documents/PukulEnam/brain-ctc-seg/Paper/jama/Supplemental_Material_Updated_CI.docx"
    save_path = "/Users/aliffandy/Documents/PukulEnam/brain-ctc-seg/Paper/jama/Supplemental_Material_FINAL.docx"
    
    if not os.path.exists(doc_path):
        print(f"File {doc_path} tidak ditemukan.")
        return

    doc = docx.Document(doc_path)
    
    # Mapping of which figure goes where
    figures_to_insert = [
        "Fig1_3Brain_Heatmap_Journal.png",
        "Fig2_Segmentation_Grid.png",
        "Fig3_Aggregated_ROC.png",
        "Fig4_Multi_CT_List.png",
        "Fig5_Mod_Seg_SE2_Architecture.png",
        "Fig6_Combined_Training_Pipeline.png"
    ]
    
    fig_idx = 0
    base_fig_dir = "/Users/aliffandy/Documents/PukulEnam/brain-ctc-seg/Paper/Journal_Figures/"
    
    # Cari paragraf yang berbunyi "Supplemental Figure X"
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("Supplemental Figure") and fig_idx < len(figures_to_insert):
            img_path = os.path.join(base_fig_dir, figures_to_insert[fig_idx])
            
            if os.path.exists(img_path):
                # Insert gambar setelah paragraf ini
                # Sayangnya docx tidak punya fungsi mudah untuk 'insert_after', 
                # jadi kita masukkan ke paragraph baru yang di-insert sebelum paragraf *berikutnya*
                if i + 1 < len(doc.paragraphs):
                    new_p = doc.paragraphs[i+1].insert_paragraph_before()
                    run = new_p.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                    print(f"✅ Berhasil menyisipkan {figures_to_insert[fig_idx]} ke bawah {text}")
                fig_idx += 1
            else:
                print(f"⚠️ Gambar tidak ditemukan: {img_path}")

    doc.save(save_path)
    print(f"\nSelesai! File disimpan sebagai: {save_path}")

if __name__ == "__main__":
    insert_figures_to_supplemental()
